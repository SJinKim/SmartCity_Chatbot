from US1_loadQA_AzureChat import qa_chain,load_file
from US1_RetrievalQA_AzureChat import re_qa_chain
from US6_bescheidtemplate import bescheidTemplate
from US5_gutachtentemplate import gutachtentemplate
import docx
import aspose.pdf as ap
import yaml

with open('config.yaml', 'r') as f:
    config = yaml.safe_load(f)

print(f'current sachverhalt path: {config["current_sv_path"]}')

def write_path_to(key, item):
    with open('config.yaml', 'r') as file:
        config = yaml.safe_load(file)
        config[f'{key}'] = f'{item}'
    print('path was modified in yaml file.')

test = True

if test:
    #Dieser ist ein zu verarbeitender Sachverhalt, der später durch den url-link vom frontend ersetzt werden sollte.
    gefragterSachverhalt = load_file("../input_docs/Sachverhalt2.docx")[0].page_content
    # Systempfad, wo die Docx-Datei vom Gutachten geschrieben und gespeichert werden soll
    gutachten_path= "../output_docs/Gutachten.docx"
    bescheid_path= "../output_docs/Beischeid.docx"
    # Systempfad, wo die Docx-Datei vom Bescheid geschrieben und gespeichert werden soll
    bescheid_path= "../output_docs/Beischeid.docx"
else:
    gefragterSachverhalt = load_file(config["current_sv_path"][0].page_content)


def strToDocx(resource: str, output_path:str) -> None:
    """ 
    Diese Funktion schreibt den Text vom Typ String in eine Docx-Datei.
    """
    doc = docx.Document() 
    doc.add_paragraph().add_run(resource) 
    doc.save(output_path)
    
def strToPdf(resource: str, output_path:str) -> None:
    """
    Diese Funktion transformiert den Text vom Typ String in deine PDF-Datei.

    Args:
        resource (_type_): _description_
        output_path (_type_): _description_
    """
    #init doc obj
    document = ap.Document()
    #add page
    page = document.pages.add()
    #init text frag obj
    txt_fragment = ap.text.TextFragment(resource)
    #add text frag to page
    page.paragraphs.add(txt_fragment)
    #save file
    document.save(output_path)

def erstelleGutachten(sachverhalt, gutachten_path)->str:
    """ein Gutachten wird als docx und pdf file format erstellt. Es wird im gutachten_path abgelegt.

    Args:
        sachverhalt (docx, pdf): sachverhalt file
        gutachten_path (str): path wo das gutachten abgelegt wird

    Returns:
        str: das Gutachten in str
    """
    #Nutzt funktion aus US-5 und den verarbeitenden Sachverhalt zum Erstellung einer Anfrage an das llm-model
    gutachten_query=gutachtentemplate(sachverhalt=sachverhalt)
    #Nutzt funktion aus US-1 zum Generieren einer Antwort für die Anfrage von einem Gutachten
    gutachten_response = qa_chain(query=gutachten_query)
    #schreibt die Antwort vom llm-model in die Datei an dem gegebenen pfad
    strToDocx(resource=gutachten_response,output_path=gutachten_path)
    strToPdf(resource=gutachten_response,output_path=gutachten_path)
    return gutachten_response 

def erstelleBescheid(sachverhalt, gutachten_result, bescheid_path)->str:
    """ein Bescheid wird als docx und pdf file Format erstellt. Es wird im gutachten_path abgelegt.

    Args:
        sachverhalt (docx, pdf): Sachverhalt file
        gutachten_result (str): return vom erstelleGutachten function
        bescheid_path (_type_): wo der Bescheid abgelegt wird

    Returns:
        str: den Bescheid in str
    """
    bescheid_query = bescheidTemplate(sachverhalt=sachverhalt,gutachten_result = gutachten_result)
    bescheid_response = qa_chain(query = bescheid_query)
    strToDocx(resource=bescheid_response,output_path=bescheid_path)    
    strToPdf(resource=bescheid_response,output_path=bescheid_path)
    #return f'Bescheid wurde in {bescheid_path} abgelegt!'
    write_path_to(key='erstellt', item=True)

    return bescheid_response


#ergebnis testen
if test:
    gutachten_response = erstelleGutachten(sachverhalt=gefragterSachverhalt,speicherpfad=gutachten_path)
    erstelleBescheid(sachverhalt=gefragterSachverhalt,gutachten_result=gutachten_response,speicherpfad=bescheid_path)
else:
    pass