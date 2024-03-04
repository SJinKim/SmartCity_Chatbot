"""
    This Modul includes the functions to generate a gutachten and a bescheid
"""

import os
from datetime import datetime
import docx
import yaml

from langchain_community.vectorstores.faiss import FAISS
from langchain.chains.question_answering import load_qa_chain

from internal.us1_load_data import init_embeddings
from internal.us3_sacherverhalt import init_llm, load_document, split_documents
from internal.us5_gutachten_template import gutachten_template
from internal.us6_bescheid_template import bescheid_template
# from internal.utils import safe_sachverhalt


llm_client = init_llm()
embeddings = init_embeddings()


def write_path_to(key, item):
    """Saves a key-value pair to a configuration YAML file.

    Args:
        key (str): The key to store the value under.
        item (Any): The value to save.

    Returns:    None
    """
    path = os.path.join(os.path.dirname(__file__), "config.yaml")
    with open(path, encoding="utf-8") as file:
        config = yaml.safe_load(file)

    config[key] = item

    with open(path, "w", encoding="utf-8") as file:
        yaml.dump(config, file)


def add_to_path(key: str, item: str) -> tuple:
    """
        add value to config file
    Args:
        key (str): key of pair to be changed
        item (str): new value

    Returns:
        tuple: updated value
    """
    with open("./internal/config.yaml", "r", encoding="utf-8") as f:
        config = yaml.safe_load(f)
    # print('start writing into yaml')
    current_chat_hist = config[key]

    current_chat_hist.append(item)
    tmp_new = current_chat_hist

    tmp_new_tuple = tuple(tmp_new)
    config[key] = tmp_new

    with open("./internal/config.yaml", "w", encoding="utf-8") as file:
        yaml.dump(config, file)

    return tmp_new_tuple


def get_value_from_config(key: str):
    """
        gets value from config with passed key
    Args:
        key (str): key

    Returns:
        str: value
    """
    with open("./internal/config.yaml", "r", encoding="utf-8") as f:
        config = yaml.safe_load(f)

    return config[key]


def str_to_docx(resource: str, output_path: str) -> None:
    """This function writes the text from a string into a Docx file.

    Args:
        resource (str): The input text as a string.
        output_path (str): The path where the Docx file will be saved.

    Returns:    None
    """
    doc = docx.Document()
    doc.add_paragraph().add_run(resource)
    doc.save(output_path)


# Verfeinerungstemplate für Gutachten- & Bescheidgenerierung
def erstelle_gutachten(sachverhalt, gutachten_path) -> str:
    """Generates an expert opinion (Gutachten) as a .docx file,
    and saves it to the specified path. This function leverages
    a pre-trained question answering chain and a pre-built indexes
    of expert opinions ("gutachten_index") to generate a new expert
    opinion based on a provided case (Sachverhalt) file.

    Args:
        sachverhalt (docx, pdf): Case file (Sachverhalt)
        gutachten_path (str): The path where the expert opinion (Gutachten) will be saved.

    Returns:
        str: The generated expert opinion (Gutachten) as a string.
    """
    gutachten_query = gutachten_template(sachverhalt=sachverhalt)

    chain = load_qa_chain(llm=llm_client, chain_type="stuff", verbose=True)
    gutachten_index = os.path.join(os.path.dirname(__file__), "gutachten_index")
    loaded_gutachten = FAISS.load_local(
        folder_path=gutachten_index, embeddings=embeddings, index_name="gutachten_index"
    )

    gutachten_response = chain.run(
        input_documents=sachverhalt, question=gutachten_query, context=loaded_gutachten
    )
    str_to_docx(resource=gutachten_response, output_path=gutachten_path)
    return gutachten_response


def erstelle_bescheid(sachverhalt, gutachten_result, bescheid_path) -> str:
    """Generates an official notice (Bescheid) as a .docx file, and saves it to the
    specified path. This function leverages a pre-trained question answering chain and
    a pre-built indexes of official notices ("bescheide_index") to generate a new
    official notice based on the generated expert opinion (Gutachten) of the provided
    case (Sachverhalt) file.

    Args:
        sachverhalt (docx, pdf): Case file (Sachverhalt)
        gutachten_result (str): The result of the `erstelleGutachten` function.
        bescheid_path (_type_): The path where the official notice (Bescheid) will be saved.

    Returns:
        str: The generated official notice (Bescheid) as a string.
    """
    bescheid_query = bescheid_template(
        sachverhalt=sachverhalt, pruefungsergebnis=gutachten_result
    )

    chain = load_qa_chain(llm=llm_client, chain_type="stuff", verbose=True)
    bescheide_index = os.path.join(os.path.dirname(__file__), "bescheide_index")
    loaded_bescheide = FAISS.load_local(
        folder_path=bescheide_index, embeddings=embeddings, index_name="bescheide_index"
    )
    similar_bescheid = loaded_bescheide.max_marginal_relevance_search(gutachten_result)

    bescheid_response = chain.run(
        input_documents=similar_bescheid, question=bescheid_query, context=sachverhalt
    )
    str_to_docx(resource=bescheid_response, output_path=bescheid_path)
    return bescheid_response


def __generated_docs_to_index(gutachten_doc, bescheid_doc):
    """
    Indexes the generated expert opinion (Gutachten) and official
    notice (Bescheid) docs by splitting them, adding them and saving them
    in the respective existing VectoreStores (gutachten_index & bescheide_index).

    Args:
        gutachten_doc (_type_): generated gutachten
        bescheid_doc (_type_): generated bescheid
    """
    split_gutachten = split_documents(gutachten_doc)
    split_bescheid = split_documents(bescheid_doc)

    gutachten_index = os.path.join(os.path.dirname(__file__), "gutachten_index")
    loaded_gutachten = FAISS.load_local(
        folder_path=gutachten_index, embeddings=embeddings, index_name="gutachten_index"
    )
    bescheide_index = os.path.join(os.path.dirname(__file__), "bescheide_index")
    loaded_bescheide = FAISS.load_local(
        folder_path=bescheide_index, embeddings=embeddings, index_name="bescheide_index"
    )

    gutachten_ids = loaded_gutachten.index_to_docstore_id
    bescheide_ids = loaded_bescheide.index_to_docstore_id
    gutachen_vorher = len(gutachten_ids)
    bescheid_vorher = len(bescheide_ids)
    loaded_gutachten.add_documents(split_gutachten)
    loaded_bescheide.add_documents(split_bescheid)
    print(
        f"Eingebettete Vektoren im < gutachten_index >: {gutachen_vorher} -> {len(gutachten_ids)}"
    )
    print(
        f"Eingebettete Vektoren im < bescheide_index >: {bescheid_vorher} -> {len(bescheide_ids)}"
    )
    loaded_gutachten.save_local(
        folder_path=gutachten_index, index_name="gutachten_index"
    )
    loaded_bescheide.save_local(
        folder_path=bescheide_index, index_name="bescheide_index"
    )

def safe_sachverhalt(file_path: str) -> None:
    sachverhalt = load_document(file_path)[0].page_content
    write_path_to(key='sachverhalt', item=sachverhalt)

# Für Fast-API in main.py
def erstelle_bescheid_background(file_path: str):
    """Asynchronously generates expert opinion (Gutachten) and official notice (Bescheid)
    documents from a case file (Sachverhalt). Timestamps are used to ensure unique file names.
    The generated documents are split into smaller chunks, and added to the existing
    VectorStores for Gutachten and Bescheide.

    Args:
        filePath (str): The path to the document containing the case file (Sachverhalt).

    Returns:    None
    """
    file_name = os.path.basename(file_path)
    timestamp = datetime.now().strftime("%Y-%m-%d")
    counter = 1

    while True:
        gutachten_path = (
            f"./Gutachten_docs/[Gutachten] {file_name}_{timestamp} [{counter}].docx"
        )
        bescheid_path = (
            f"./Bescheide_docs/[Bescheid] {file_name}_{timestamp} [{counter}].docx"
        )

        if not os.path.exists(gutachten_path) and not os.path.exists(bescheid_path):
            break

        counter += 1
    
    safe_sachverhalt(file_path)

    gutachten = erstelle_gutachten(
        sachverhalt=load_document(file_path), gutachten_path=gutachten_path
    )
    message_str = erstelle_bescheid(
        sachverhalt=load_document(file_path),
        gutachten_result=gutachten,
        bescheid_path=bescheid_path,
    )

    __generated_docs_to_index(
        load_document(gutachten_path), load_document(bescheid_path)
    )

    write_path_to(key="message_str", item=message_str)
    write_path_to(key="erstellt", item=True)
