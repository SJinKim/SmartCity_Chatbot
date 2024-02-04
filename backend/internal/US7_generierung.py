from internal.US3_sacherverhalt import execute_qa_chain, load_document
from internal.US5_gutachtentemplate import gutachtentemplate
from internal.US6_bescheidtemplate import bescheidTemplate

import os
import yaml
import docx
import PyPDF2



def write_path_to(key, item):
    with open('./internal/config.yaml') as file:
        config = yaml.safe_load(file)    
    
    config[key] = item

    with open('./internal/config.yaml', 'w') as file:
        yaml.dump(config, file)


def strToDocx(resource: str, output_path:str) -> None:
    """ 
    Diese Funktion schreibt den Text vom Typ String in eine Docx-Datei.
    """
    doc = docx.Document() 
    doc.add_paragraph().add_run(resource) 
    doc.save(output_path)
    
def strToPdf(resource: str, output_path: str) -> None:
    """
    This function transforms the text from a string into a PDF file.

    Args:
        resource (str): The input text as a string.
        output_path (str): The path where the PDF file will be saved.
    """
    pdf_writer = PyPDF2.PdfWriter()
    pdf_writer.addPage()
    pdf_writer.write(resource)
    
    with open(output_path, 'wb') as output_pdf:
        pdf_writer.write(output_pdf)


def erstelleGutachten(sachverhalt, gutachten_path)->str:
    """ein Gutachten wird als docx und pdf file format erstellt. Es wird im gutachten_path abgelegt.

    Args:
        sachverhalt (docx, pdf): sachverhalt file
        gutachten_path (str): path wo das gutachten abgelegt wird

    Returns:
        str: das Gutachten in str
    """
    #Nutzt funktion aus US-5 und den verarbeitenden Sachverhalt zum Erstellung einer Anfrage an das llm-model
    gutachten_query=gutachtentemplate(sachverhalt=sachverhalt)
    #Nutzt funktion aus US-1 zum Generieren einer Antwort für die Anfrage von einem Gutachten
    gutachten_response = execute_qa_chain(message=gutachten_query)
    #schreibt die Antwort vom llm-model in die Datei an dem gegebenen pfad
    strToDocx(resource=gutachten_response,output_path=gutachten_path)
    #strToPdf(resource=gutachten_response,output_path=gutachten_path)
    return gutachten_response 

def erstelleBescheid(sachverhalt, gutachten_result, bescheid_path)->str:
    """ein Bescheid wird als docx und pdf file Format erstellt. Es wird im gutachten_path abgelegt.

    Args:
        sachverhalt (docx, pdf): Sachverhalt file
        gutachten_result (str): return vom erstelleGutachten function
        bescheid_path (_type_): wo der Bescheid abgelegt wird

    Returns:
        str: den Bescheid in str
    """
    bescheid_query = bescheidTemplate(sachverhalt=sachverhalt,prüfungsergebnis=gutachten_result)
    bescheid_response = execute_qa_chain(message=bescheid_query)
    strToDocx(resource=bescheid_response,output_path=bescheid_path)    
    #strToPdf(resource=bescheid_response,output_path=bescheid_path)
    return bescheid_response

def erstelleBescheidBackground(filePath: str):
    file_name = os.path.basename(filePath)
    gutachten_name = "Gutachten." + os.path.splitext(file_name)[0] + ".docx"
    bescheid_name = "Bescheid." + os.path.splitext(file_name)[0] + ".docx"

    gutachten = erstelleGutachten(sachverhalt=load_document(filePath), gutachten_path=f"./output_docs/{gutachten_name}")
    message_str = erstelleBescheid(sachverhalt=load_document(filePath), gutachten_result=gutachten, bescheid_path=f"./output_docs/{bescheid_name}")
    write_path_to(key='message_str', item=message_str)
    write_path_to(key='erstellt', item=True)

